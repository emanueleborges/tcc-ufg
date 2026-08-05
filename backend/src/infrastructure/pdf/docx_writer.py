"""Geração de documentos DOCX a partir de texto/markdown simples."""

from __future__ import annotations

import re
from pathlib import Path


class DocxWriter:
    """Converte texto/markdown simples em .docx via python-docx."""

    def text_to_docx(self, text: str, output_path: Path) -> Path:
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError as exc:
            raise RuntimeError(
                "Instale python-docx para gerar DOCX: pip install python-docx"
            ) from exc

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        for raw_line in (text or "").splitlines():
            line = raw_line.rstrip()
            if not line.strip():
                document.add_paragraph("")
                continue
            if line.startswith("# "):
                document.add_heading(line[2:].strip(), level=1)
            elif line.startswith("## "):
                document.add_heading(line[3:].strip(), level=2)
            elif line.startswith("### "):
                document.add_heading(line[4:].strip(), level=3)
            elif re.match(r"^[-*]\s+", line):
                document.add_paragraph(
                    re.sub(r"^[-*]\s+", "", line),
                    style="List Bullet",
                )
            else:
                document.add_paragraph(_strip_md_inline(line))

        document.save(str(output_path))
        return output_path


def _strip_md_inline(text: str) -> str:
    cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    cleaned = re.sub(r"\*(.+?)\*", r"\1", cleaned)
    cleaned = re.sub(r"`(.+?)`", r"\1", cleaned)
    return cleaned
